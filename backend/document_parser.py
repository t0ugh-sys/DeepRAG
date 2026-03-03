"""
深度文档理解模块
支持 OCR、表格提取、多格式文档解析
"""
import os
from typing import Dict, List, Any, Tuple
from dataclasses import dataclass
import logging

logger = logging.getLogger("rag")


@dataclass
class ParsedDocument:
    """解析后的文档结构"""
    text: str
    tables: List[Dict[str, Any]]
    images: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    page_info: List[Dict[str, Any]]  # 页码信息，用于精确引用


class DocumentParser:
    """统一文档解析器"""
    
    def __init__(self, enable_ocr: bool = False):
        """
        初始化文档解析器
        
        Args:
            enable_ocr: 是否启用 OCR（需要安装 Tesseract）
        """
        self.enable_ocr = enable_ocr
        
    def parse(self, file_path: str) -> ParsedDocument:
        """
        解析文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            ParsedDocument: 解析结果
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._parse_word(file_path)
        elif ext in ['.xlsx', '.xls']:
            return self._parse_excel(file_path)
        elif ext in ['.txt', '.md']:
            return self._parse_text(file_path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            return self._parse_image(file_path)
        else:
            # 回退到简单文本读取
            return self._parse_text(file_path)

    @staticmethod
    def _build_pdf_text_and_page_info(pages: List[Dict[str, Any]]) -> Tuple[str, List[Dict[str, Any]]]:
        """
        根据每页提取结果构建最终 text 与 page_info，保证 char_start/char_end 与最终文本严格一致。

        pages 每项格式：
        - page: int
        - text: str（该页纯文本，可为空）
        - segments: List[str]（该页附加片段，如表格 markdown，按出现顺序）
        """
        parts: List[str] = []
        page_info: List[Dict[str, Any]] = []
        cursor = 0

        for page in pages:
            page_num = int(page["page"])
            page_text = str(page.get("text") or "")
            segments = list(page.get("segments") or [])

            page_segments: List[str] = []
            if page_text.strip():
                page_segments.append(page_text)
            page_segments.extend([s for s in segments if str(s).strip()])

            if not page_segments:
                continue

            page_start: int | None = None
            for seg in page_segments:
                if parts:
                    parts.append("\n\n")
                    cursor += 2
                if page_start is None:
                    page_start = cursor
                parts.append(seg)
                cursor += len(seg)

            page_end = cursor
            page_info.append(
                {
                    "page": page_num,
                    "text": page_text,
                    "char_start": page_start if page_start is not None else page_end,
                    "char_end": page_end,
                }
            )

        return "".join(parts), page_info
    
    def _parse_pdf(self, file_path: str) -> ParsedDocument:
        """解析 PDF 文件，提取文本、表格和图片"""
        try:
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber 未安装，回退到基础 PDF 解析")
            return self._parse_pdf_basic(file_path)
        
        tables = []
        images = []
        pages: List[Dict[str, Any]] = []
        
        with pdfplumber.open(file_path) as pdf:
            pages_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, start=1):
                page_record: Dict[str, Any] = {"page": page_num, "text": "", "segments": []}

                # 提取文本
                page_text = page.extract_text() or ""
                page_record["text"] = page_text
                
                # 提取表格
                page_tables = page.extract_tables()
                for table_idx, table in enumerate(page_tables):
                    if table:
                        # 转换表格为 Markdown 格式
                        table_md = self._table_to_markdown(table)
                        table_number = len(tables) + 1
                        tables.append({
                            "page": page_num,
                            "index": table_idx,
                            "data": table,
                            "markdown": table_md
                        })
                        page_record["segments"].append(f"\n[表格 {table_number}]\n{table_md}\n")
                
                # 提取图片（如果启用 OCR）
                if self.enable_ocr:
                    page_images = page.images
                    for img_idx, img in enumerate(page_images):
                        images.append({
                            "page": page_num,
                            "index": img_idx,
                            "bbox": (img.get("x0"), img.get("top"), img.get("x1"), img.get("bottom"))
                        })

                pages.append(page_record)

        text, page_info = self._build_pdf_text_and_page_info(pages)
        
        return ParsedDocument(
            text=text,
            tables=tables,
            images=images,
            metadata={"pages": pages_count, "format": "pdf"},
            page_info=page_info
        )
    
    def _parse_pdf_basic(self, file_path: str) -> ParsedDocument:
        """基础 PDF 解析（回退方案）"""
        from pdfminer.high_level import extract_text
        text = extract_text(file_path)
        return ParsedDocument(
            text=text,
            tables=[],
            images=[],
            metadata={"format": "pdf", "parser": "pdfminer"},
            page_info=[]
        )
    
    def _parse_word(self, file_path: str) -> ParsedDocument:
        """解析 Word 文档"""
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx 未安装，无法解析 Word 文档")
            return ParsedDocument(text="", tables=[], images=[], metadata={}, page_info=[])
        
        doc = Document(file_path)
        text_parts = []
        tables = []
        
        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                table_data.append([cell.text for cell in row.cells])
            
            table_md = self._table_to_markdown(table_data)
            tables.append({
                "index": table_idx,
                "data": table_data,
                "markdown": table_md
            })
            text_parts.append(f"\n[表格 {table_idx + 1}]\n{table_md}\n")
        
        return ParsedDocument(
            text="\n\n".join(text_parts),
            tables=tables,
            images=[],
            metadata={"format": "docx", "paragraphs": len(doc.paragraphs)},
            page_info=[]
        )
    
    def _parse_excel(self, file_path: str) -> ParsedDocument:
        """解析 Excel 文件"""
        try:
            from openpyxl import load_workbook
        except ImportError:
            logger.warning("openpyxl 未安装，无法解析 Excel 文件")
            return ParsedDocument(text="", tables=[], images=[], metadata={}, page_info=[])
        
        wb = load_workbook(file_path, data_only=True)
        text_parts = []
        tables = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"## {sheet_name}\n")
            
            # 读取所有数据
            table_data = []
            for row in sheet.iter_rows(values_only=True):
                if any(cell is not None for cell in row):
                    table_data.append([str(cell) if cell is not None else "" for cell in row])
            
            if table_data:
                table_md = self._table_to_markdown(table_data)
                tables.append({
                    "sheet": sheet_name,
                    "data": table_data,
                    "markdown": table_md
                })
                text_parts.append(table_md)
        
        return ParsedDocument(
            text="\n\n".join(text_parts),
            tables=tables,
            images=[],
            metadata={"format": "xlsx", "sheets": len(wb.sheetnames)},
            page_info=[]
        )
    
    def _parse_text(self, file_path: str) -> ParsedDocument:
        """解析纯文本文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return ParsedDocument(
            text=text,
            tables=[],
            images=[],
            metadata={"format": "text"},
            page_info=[]
        )
    
    def _parse_image(self, file_path: str) -> ParsedDocument:
        """解析图片文件（OCR）"""
        if not self.enable_ocr:
            return ParsedDocument(
                text=f"[图片: {os.path.basename(file_path)}]",
                tables=[],
                images=[{"path": file_path}],
                metadata={"format": "image"},
                page_info=[]
            )
        
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            logger.warning("pytesseract 或 Pillow 未安装，无法进行 OCR")
            return ParsedDocument(
                text=f"[图片: {os.path.basename(file_path)}]",
                tables=[],
                images=[{"path": file_path}],
                metadata={"format": "image"},
                page_info=[]
            )
        
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            return ParsedDocument(
                text=text,
                tables=[],
                images=[{"path": file_path, "ocr": True}],
                metadata={"format": "image", "ocr": True},
                page_info=[]
            )
        except Exception as e:
            logger.error(f"OCR 失败: {e}")
            return ParsedDocument(
                text=f"[图片: {os.path.basename(file_path)}]",
                tables=[],
                images=[{"path": file_path}],
                metadata={"format": "image"},
                page_info=[]
            )
    
    @staticmethod
    def _table_to_markdown(table: List[List[str]]) -> str:
        """将表格转换为 Markdown 格式"""
        if not table or not table[0]:
            return ""
        
        lines = []
        # 表头
        lines.append("| " + " | ".join(str(cell) for cell in table[0]) + " |")
        # 分隔线
        lines.append("| " + " | ".join("---" for _ in table[0]) + " |")
        # 数据行
        for row in table[1:]:
            lines.append("| " + " | ".join(str(cell) for cell in row) + " |")
        
        return "\n".join(lines)
